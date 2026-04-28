"""
tools/document_parser.py
────────────────────────
Document parsing utilities for the Ingestion pipeline stage.

Supports:
    .pdf   – pdfplumber (already in dependencies)
    .docx  – python-docx (already in dependencies)
    .xlsx  – openpyxl  (optional; falls back gracefully)
    .xls   – openpyxl
    .csv   – stdlib csv
    .txt   – plain read
    .md    – plain read

Usage::

    from app.tools.document_parser import parse_document

    text = parse_document("/path/to/spec.pdf")
    text = parse_document("/path/to/requirements.docx")
"""

from __future__ import annotations

import csv
import logging
from pathlib import Path
from typing import Callable

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────────────────────────────────────


def _read_text_file(file_path: str | Path) -> str:
    """Read a plain-text / Markdown / CSV file and return its contents."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix == ".csv":
        return _parse_csv(path)

    content = path.read_text(encoding="utf-8", errors="replace")
    logger.info("Parsed text file %r: %d chars", path.name, len(content))
    return content


def _parse_csv(path: Path) -> str:
    """Convert CSV rows → newline-delimited pipe-separated text."""
    rows: list[str] = []
    with path.open(newline="", encoding="utf-8", errors="replace") as fh:
        reader = csv.reader(fh)
        for row in reader:
            line = " | ".join(cell.strip() for cell in row if cell.strip())
            if line:
                rows.append(line)
    content = "\n".join(rows)
    logger.info("Parsed CSV %r: %d rows, %d chars", path.name, len(rows), len(content))
    return content


# ─────────────────────────────────────────────────────────────────────────────
# Public per-format parsers
# ─────────────────────────────────────────────────────────────────────────────


def parse_pdf(file_path: str | Path) -> str:
    """
    Extract text from a PDF file using *pdfplumber*.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Concatenated text of all pages, separated by a page-break marker.

    Raises:
        FileNotFoundError: If the file does not exist.
        ImportError: If pdfplumber is not installed.
    """
    try:
        import pdfplumber  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError(
            "pdfplumber is not installed. Run: uv add pdfplumber"
        ) from exc

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    pages_text: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            raw = page.extract_text() or ""
            text = raw.strip()
            if text:
                pages_text.append(f"[Page {page_num}]\n{text}")

    result = "\n\n".join(pages_text)
    logger.info(
        "Parsed PDF %r: %d pages with text, %d chars",
        path.name,
        len(pages_text),
        len(result),
    )
    return result


def parse_docx(file_path: str | Path) -> str:
    """
    Extract text from a DOCX (Word) file using *python-docx*.

    Extracts both body paragraphs and table cell content so that
    requirement tables in Word documents are captured.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Newline-joined paragraph text.

    Raises:
        FileNotFoundError: If the file does not exist.
        ImportError: If python-docx is not installed.
    """
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError(
            "python-docx is not installed. Run: uv add python-docx"
        ) from exc

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    doc = docx.Document(str(path))
    parts: list[str] = []

    # --- body paragraphs ---
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading hierarchy with a prefix so the LLM can use it
            if para.style and "Heading" in para.style.name:
                level = "".join(filter(str.isdigit, para.style.name)) or "1"
                text = f"{'#' * int(level)} {text}"
            parts.append(text)

    # --- table cells ---
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    result = "\n".join(parts)
    logger.info(
        "Parsed DOCX %r: %d text blocks, %d chars",
        path.name,
        len(parts),
        len(result),
    )
    return result


def parse_excel(file_path: str | Path) -> str:
    """
    Extract text from an Excel workbook (.xlsx / .xls) using *openpyxl*.

    Each worksheet is output as a section with a ``[Sheet: name]`` header,
    and each row is rendered as pipe-separated values.

    Args:
        file_path: Path to the Excel file.

    Returns:
        Multi-sheet text representation.

    Raises:
        FileNotFoundError: If the file does not exist.
        ImportError: If openpyxl is not installed.
    """
    try:
        import openpyxl  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError("openpyxl is not installed. Run: uv add openpyxl") from exc

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Excel file not found: {path}")

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    sheet_blocks: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [
                str(cell).strip()
                for cell in row
                if cell is not None and str(cell).strip()
            ]
            if cells:
                rows_text.append(" | ".join(cells))

        if rows_text:
            block = f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text)
            sheet_blocks.append(block)

    wb.close()
    result = "\n\n".join(sheet_blocks)
    logger.info(
        "Parsed Excel %r: %d sheet(s), %d chars",
        path.name,
        len(sheet_blocks),
        len(result),
    )
    return result


# ─────────────────────────────────────────────────────────────────────────────
# Auto-dispatch entry point
# ─────────────────────────────────────────────────────────────────────────────

_EXTENSION_MAP: dict[str, Callable[[str | Path], str]] = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".doc": parse_docx,
    ".xlsx": parse_excel,
    ".xls": parse_excel,
    ".txt": _read_text_file,
    ".md": _read_text_file,
    ".csv": _read_text_file,
}


def parse_document(file_path: str | Path) -> str:
    """
    Auto-detect the file format and return its text content.

    Supported formats: .pdf, .docx, .doc, .xlsx, .xls, .txt, .md, .csv

    Args:
        file_path: Path to the document.

    Returns:
        Extracted text, ready for chunking and LLM analysis.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file format is not supported.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    suffix = path.suffix.lower()
    parser = _EXTENSION_MAP.get(suffix)

    if parser is None:
        supported = sorted(_EXTENSION_MAP.keys())
        raise ValueError(f"Unsupported file format {suffix!r}. Supported: {supported}")

    logger.debug("Parsing document %r with format %r", path.name, suffix)
    return parser(file_path)


def supported_extensions() -> list[str]:
    """Return the list of supported file extensions (with leading dot)."""
    return sorted(_EXTENSION_MAP.keys())
