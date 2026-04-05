from __future__ import annotations

"""
services/docx_builder.py
────────────────────────
DocxReportBuilder: Builds a formatted Microsoft Word DOCX report
from Auto-AT pipeline results.

Usage::

    builder = DocxReportBuilder()
    builder.add_title_page("MyApp SRS.pdf", "abc12345", "2024-01-15 10:00:00 UTC", "completed")
    builder.add_executive_summary("All tests passed.", {"Pass Rate": "100%", ...})
    builder.add_test_cases_table(test_cases, exec_results)
    docx_bytes = builder.build()
"""

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ─────────────────────────────────────────────────────────────────────────────
# Colour palette
# ─────────────────────────────────────────────────────────────────────────────

_DARK_BLUE = RGBColor(0x1E, 0x3A, 0x5F)  # #1e3a5f  — headings
_ACCENT = RGBColor(0x3B, 0x82, 0xF6)  # #3b82f6  — accent / links
_PASS = RGBColor(0x22, 0xC5, 0x5E)  # #22c55e  — passed
_FAIL = RGBColor(0xEF, 0x44, 0x44)  # #ef4444  — failed
_SKIP = RGBColor(0xF5, 0x9E, 0x0B)  # #f59e0b  — skipped
_ERROR_CLR = RGBColor(0xF9, 0x73, 0x16)  # #f97316  — error
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_LIGHT_BG = "F8FAFC"  # alternating row shade (no #)
_HEADER_BG = "1E3A5F"  # table header background


# ─────────────────────────────────────────────────────────────────────────────
# XML helpers
# ─────────────────────────────────────────────────────────────────────────────


def _set_cell_bg(cell: Any, hex_color: str) -> None:
    """Set the background fill colour of a table cell.

    Args:
        cell:      A ``docx`` table cell object.
        hex_color: Six-character hex colour string **without** leading ``#``.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell: Any, *, top: bool = False, bottom: bool = True) -> None:
    """Add a simple 1pt border to the top and/or bottom of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, active in (("top", top), ("bottom", bottom)):
        if active:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "E2E8F0")
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_text(
    cell: Any,
    text: str,
    bold: bool = False,
    font_size: int = 10,
    color: RGBColor | None = None,
    italic: bool = False,
) -> None:
    """Clear a cell and write plain formatted text into it.

    Args:
        cell:      Target table cell.
        text:      String to write.
        bold:      Whether to apply bold formatting.
        font_size: Point size of the font.
        color:     Optional :class:`~docx.shared.RGBColor` for the text.
        italic:    Whether to apply italic formatting.
    """
    cell.text = ""
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    if color is not None:
        run.font.color.rgb = color


def _status_color(status: str) -> RGBColor:
    """Return the :class:`~docx.shared.RGBColor` associated with an execution status.

    Args:
        status: One of ``"passed"``, ``"failed"``, ``"skipped"``, ``"error"``.

    Returns:
        Appropriate :class:`~docx.shared.RGBColor` instance.
    """
    mapping = {
        "passed": _PASS,
        "failed": _FAIL,
        "skipped": _SKIP,
        "error": _ERROR_CLR,
    }
    return mapping.get(status.lower(), RGBColor(0x64, 0x74, 0x8B))


# ─────────────────────────────────────────────────────────────────────────────
# DocxReportBuilder
# ─────────────────────────────────────────────────────────────────────────────


class DocxReportBuilder:
    """Build a formatted Microsoft Word ``.docx`` report from Auto-AT pipeline data.

    Sections are added in order via explicit ``add_*`` method calls.
    Call :meth:`build` at the end to serialise the document to bytes.

    Example::

        builder = DocxReportBuilder()
        builder.add_title_page("SRS.pdf", "abc1234", "2024-01-15 10:00 UTC", "completed")
        builder.add_executive_summary("All systems go.", {"Pass Rate": "95.0%"})
        builder.add_test_cases_table(test_cases, exec_results)
        docx_bytes = builder.build()
    """

    def __init__(self) -> None:
        self.doc = Document()
        self._setup_styles()

    # ── Styles ────────────────────────────────────────────────────────────────

    def _setup_styles(self) -> None:
        """Configure built-in styles for the document.

        Sets ``Normal`` to Calibri 11 pt, and tunes Heading 1 / 2 / 3 to use
        the dark-blue brand colour at 18 / 14 / 12 pt respectively.
        """
        styles = self.doc.styles

        # Normal
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        # Heading 1
        h1 = styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = _DARK_BLUE

        # Heading 2
        h2 = styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = _DARK_BLUE

        # Heading 3
        h3 = styles["Heading 3"]
        h3.font.name = "Calibri"
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = _ACCENT

        # Page margins (1 inch all around)
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

    # ── Title Page ────────────────────────────────────────────────────────────

    def add_title_page(
        self,
        document_name: str,
        run_id: str,
        generated_at: str,
        status: str,
    ) -> None:
        """Add a title page with document name, subtitle, and run metadata.

        Args:
            document_name: Original filename of the requirements document.
            run_id:        UUID string of the pipeline run.
            generated_at:  Human-readable datetime string (UTC).
            status:        Final pipeline status (e.g. ``"completed"``).
        """
        doc = self.doc

        # Vertical spacing before title
        for _ in range(4):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)

        # Document / project name
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(document_name)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = _DARK_BLUE
        title_run.font.name = "Calibri"

        # Report subtitle
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run("Automated Test Report")
        sub_run.font.size = Pt(16)
        sub_run.font.color.rgb = _ACCENT
        sub_run.font.name = "Calibri"

        doc.add_paragraph()  # spacer

        # Divider line via a 1-cell, 0-border table with bottom border
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.rows[0].cells[0]
        cell.text = ""
        _set_cell_bg(cell, "1E3A5F")
        cell_para = cell.paragraphs[0]
        cell_para.paragraph_format.space_before = Pt(2)
        cell_para.paragraph_format.space_after = Pt(2)

        doc.add_paragraph()  # spacer

        # Metadata block
        meta_lines = [
            ("Generated", generated_at),
            ("Pipeline Run ID", run_id),
            ("Status", status.title()),
            ("Generated by", "Auto-AT"),
        ]
        for label, value in meta_lines:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lbl_run = meta_para.add_run(f"{label}:  ")
            lbl_run.font.size = Pt(11)
            lbl_run.font.bold = True
            lbl_run.font.color.rgb = _DARK_BLUE
            lbl_run.font.name = "Calibri"
            val_run = meta_para.add_run(value)
            val_run.font.size = Pt(11)
            val_run.font.name = "Calibri"

        # Page break
        doc.add_page_break()

    # ── Executive Summary ─────────────────────────────────────────────────────

    def add_executive_summary(self, summary: str, metrics: dict[str, str]) -> None:
        """Add the Executive Summary section with a metrics table.

        Args:
            summary: Prose summary paragraph.
            metrics: Ordered mapping of metric label → formatted value
                     (e.g. ``{"Pass Rate": "95.0%", "Coverage": "88.0%"}``).
        """
        doc = self.doc
        doc.add_heading("Executive Summary", level=1)

        if summary:
            p = doc.add_paragraph(summary)
            p.paragraph_format.space_after = Pt(10)

        if not metrics:
            return

        doc.add_paragraph()  # spacer before table

        # 2-column metrics table: label | value
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"

        # Header row
        hdr_cells = tbl.rows[0].cells
        for cell, text in zip(hdr_cells, ["Metric", "Value"]):
            _set_cell_bg(cell, _HEADER_BG)
            _cell_text(cell, text, bold=True, font_size=10, color=_WHITE)

        for i, (label, value) in enumerate(metrics.items()):
            row = tbl.add_row()
            bg = _LIGHT_BG if i % 2 == 0 else "FFFFFF"
            _set_cell_bg(row.cells[0], bg)
            _set_cell_bg(row.cells[1], bg)
            _cell_text(row.cells[0], label, bold=True, font_size=10)
            _cell_text(row.cells[1], value, font_size=10)

        doc.add_paragraph()  # spacer after table

    # ── Requirements ─────────────────────────────────────────────────────────

    def add_requirements_section(self, requirements: list[dict]) -> None:  # type: ignore[type-arg]
        """Add the Requirements section with an overview table.

        Only called when *requirements* is non-empty.

        Args:
            requirements: List of serialised :class:`RequirementItem` dicts,
                          each containing ``id``, ``title``, ``type``,
                          ``priority``, and optionally ``description``.
        """
        if not requirements:
            return

        doc = self.doc
        doc.add_heading("Requirements", level=1)

        cols = ["ID", "Title", "Type", "Priority"]
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Table Grid"

        # Header
        for cell, text in zip(tbl.rows[0].cells, cols):
            _set_cell_bg(cell, _HEADER_BG)
            _cell_text(cell, text, bold=True, font_size=10, color=_WHITE)

        for i, req in enumerate(requirements):
            row = tbl.add_row()
            bg = _LIGHT_BG if i % 2 == 0 else "FFFFFF"
            values = [
                req.get("id", ""),
                req.get("title", ""),
                req.get("type", ""),
                req.get("priority", ""),
            ]
            for cell, val in zip(row.cells, values):
                _set_cell_bg(cell, bg)
                _cell_text(cell, val, font_size=10)

        doc.add_paragraph()

    # ── Test Cases Table ──────────────────────────────────────────────────────

    def add_test_cases_table(
        self,
        test_cases: list[dict],  # type: ignore[type-arg]
        exec_results: list[dict],  # type: ignore[type-arg]
    ) -> None:
        """Add the Test Cases section with execution status joined in.

        Only called when *test_cases* is non-empty.

        Args:
            test_cases:   List of serialised :class:`TestCase` dicts.
            exec_results: List of serialised :class:`TestExecutionResult` dicts.
                          Joined to test cases on ``test_case_id == tc["id"]``.
        """
        if not test_cases:
            return

        doc = self.doc
        doc.add_heading("Test Cases", level=1)

        # Build a quick lookup: test_case_id → execution status
        status_map: dict[str, str] = {
            r["test_case_id"]: r.get("status", "not run")
            for r in exec_results
            if isinstance(r, dict) and "test_case_id" in r
        }

        cols = ["ID", "Title", "Requirement", "Type", "Category", "Priority", "Status"]
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Table Grid"

        # Header
        for cell, text in zip(tbl.rows[0].cells, cols):
            _set_cell_bg(cell, _HEADER_BG)
            _cell_text(cell, text, bold=True, font_size=9, color=_WHITE)

        for i, tc in enumerate(test_cases):
            row = tbl.add_row()
            bg = _LIGHT_BG if i % 2 == 0 else "FFFFFF"
            tc_id = tc.get("id", "")
            raw_status = status_map.get(tc_id, "not run")

            values = [
                tc_id,
                tc.get("title", ""),
                tc.get("requirement_id", ""),
                tc.get("test_type", ""),
                tc.get("category", ""),
                tc.get("priority", ""),
            ]
            for cell, val in zip(row.cells, values):
                _set_cell_bg(cell, bg)
                _cell_text(cell, val, font_size=9)

            # Status cell with colour
            status_cell = row.cells[-1]
            _set_cell_bg(status_cell, bg)
            _cell_text(
                status_cell,
                raw_status.title(),
                font_size=9,
                bold=True,
                color=_status_color(raw_status),
            )

        doc.add_paragraph()

    # ── Execution Summary ─────────────────────────────────────────────────────

    def add_execution_summary(self, summary: dict) -> None:  # type: ignore[type-arg]
        """Add the Execution Summary section.

        Only called when *summary* is non-empty.

        Args:
            summary: Serialised :class:`ExecutionSummary` dict with fields
                     ``total``, ``passed``, ``failed``, ``skipped``, ``errors``,
                     ``pass_rate``, and ``duration_seconds``.
        """
        if not summary:
            return

        doc = self.doc
        doc.add_heading("Execution Summary", level=1)

        rows_data = [
            ("Total Tests", str(summary.get("total", 0))),
            ("Passed", str(summary.get("passed", 0))),
            ("Failed", str(summary.get("failed", 0))),
            ("Skipped", str(summary.get("skipped", 0))),
            ("Errors", str(summary.get("errors", 0))),
            ("Pass Rate", f"{summary.get('pass_rate', 0.0):.1f}%"),
            (
                "Duration",
                f"{summary.get('duration_seconds', 0.0):.2f}s",
            ),
        ]

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"

        for cell, text in zip(tbl.rows[0].cells, ["Metric", "Value"]):
            _set_cell_bg(cell, _HEADER_BG)
            _cell_text(cell, text, bold=True, font_size=10, color=_WHITE)

        for i, (label, value) in enumerate(rows_data):
            row = tbl.add_row()
            bg = _LIGHT_BG if i % 2 == 0 else "FFFFFF"
            _set_cell_bg(row.cells[0], bg)
            _set_cell_bg(row.cells[1], bg)

            # Colour-code numeric values for passed/failed
            val_color: RGBColor | None = None
            if label == "Passed":
                val_color = _PASS
            elif label in ("Failed", "Errors"):
                val_color = _FAIL
            elif label == "Skipped":
                val_color = _SKIP

            _cell_text(row.cells[0], label, bold=True, font_size=10)
            _cell_text(row.cells[1], value, font_size=10, color=val_color)

        doc.add_paragraph()

    # ── Coverage Analysis ─────────────────────────────────────────────────────

    def add_coverage_section(self, coverage: dict) -> None:  # type: ignore[type-arg]
        """Add the Coverage Analysis section.

        Only called when *coverage* is non-empty.

        Args:
            coverage: Serialised :class:`PostExecutionCoverage` dict.
        """
        if not coverage:
            return

        doc = self.doc
        doc.add_heading("Coverage Analysis", level=1)

        rows_data = [
            ("Total Requirements", str(coverage.get("total_requirements", 0))),
            ("Covered Requirements", str(coverage.get("covered_requirements", 0))),
            ("Validated Requirements", str(coverage.get("validated_requirements", 0))),
            ("Coverage %", f"{coverage.get('coverage_percentage', 0.0):.1f}%"),
            ("Validation %", f"{coverage.get('validation_percentage', 0.0):.1f}%"),
        ]

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"

        for cell, text in zip(tbl.rows[0].cells, ["Metric", "Value"]):
            _set_cell_bg(cell, _HEADER_BG)
            _cell_text(cell, text, bold=True, font_size=10, color=_WHITE)

        for i, (label, value) in enumerate(rows_data):
            row = tbl.add_row()
            bg = _LIGHT_BG if i % 2 == 0 else "FFFFFF"
            _set_cell_bg(row.cells[0], bg)
            _set_cell_bg(row.cells[1], bg)
            _cell_text(row.cells[0], label, bold=True, font_size=10)
            _cell_text(row.cells[1], value, font_size=10)

        # Uncovered requirements
        uncovered: list[str] = coverage.get("uncovered_requirements") or []
        if uncovered:
            doc.add_paragraph()
            h = doc.add_heading("Uncovered Requirements", level=2)
            h.paragraph_format.space_before = Pt(4)
            for req_id in uncovered:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(req_id).font.size = Pt(10)

        doc.add_paragraph()

    # ── Root Cause Analysis ───────────────────────────────────────────────────

    def add_root_cause_section(self, root_causes: list[dict]) -> None:  # type: ignore[type-arg]
        """Add the Root Cause Analysis section.

        Only called when *root_causes* is non-empty.

        Args:
            root_causes: List of serialised :class:`RootCause` dicts, each
                         with ``test_case_id``, ``failure_pattern``,
                         ``root_cause_category``, ``severity``,
                         ``recommendation``.
        """
        if not root_causes:
            return

        doc = self.doc
        doc.add_heading("Root Cause Analysis", level=1)

        cols = ["Test Case", "Pattern", "Category", "Severity", "Recommendation"]
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Table Grid"

        for cell, text in zip(tbl.rows[0].cells, cols):
            _set_cell_bg(cell, _HEADER_BG)
            _cell_text(cell, text, bold=True, font_size=9, color=_WHITE)

        for i, rc in enumerate(root_causes):
            row = tbl.add_row()
            bg = _LIGHT_BG if i % 2 == 0 else "FFFFFF"

            values = [
                rc.get("test_case_id", ""),
                rc.get("failure_pattern", ""),
                rc.get("root_cause_category", ""),
                rc.get("severity", ""),
                rc.get("recommendation", ""),
            ]
            for j, (cell, val) in enumerate(zip(row.cells, values)):
                _set_cell_bg(cell, bg)
                # Colour-code severity cell
                color: RGBColor | None = None
                if j == 3:  # severity column
                    sev = str(val).lower()
                    color = (
                        _FAIL
                        if sev == "critical"
                        else (
                            _ERROR_CLR
                            if sev == "high"
                            else (_SKIP if sev == "medium" else None)
                        )
                    )
                _cell_text(
                    cell,
                    val,
                    font_size=9,
                    color=color,
                    bold=(j == 3 and color is not None),
                )

        doc.add_paragraph()

    # ── Recommendations & Risks ───────────────────────────────────────────────

    def add_recommendations(
        self,
        recommendations: list[str],
        risks: list[str],
    ) -> None:
        """Add the Recommendations and Risks section.

        Only called when at least one list is non-empty.

        Args:
            recommendations: List of recommendation strings.
            risks:           List of risk description strings.
        """
        doc = self.doc
        doc.add_heading("Recommendations & Risks", level=1)

        if recommendations:
            doc.add_heading("Recommendations", level=2)
            for i, rec in enumerate(recommendations, start=1):
                p = doc.add_paragraph(style="List Number")
                run = p.add_run(rec)
                run.font.size = Pt(10)

        if risks:
            doc.add_paragraph()
            doc.add_heading("Risks", level=2)
            for risk in risks:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(risk)
                run.font.size = Pt(10)
                run.font.color.rgb = _FAIL

        doc.add_paragraph()

    # ── Build ─────────────────────────────────────────────────────────────────

    def build(self) -> bytes:
        """Serialise the document to a raw bytes buffer.

        Returns:
            Raw ``.docx`` bytes suitable for streaming as an HTTP response.
        """
        buffer = BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()
